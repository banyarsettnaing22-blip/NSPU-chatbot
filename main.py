import os
import json
import pymysql
import chromadb
from typing import List, Optional
from pathlib import Path
from dotenv import load_dotenv

# --- 1. Robust .env Path Loading ---
BASE_DIR = Path(__file__).resolve().parent
load_dotenv(dotenv_path=BASE_DIR / ".env")

from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect, Query, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import OpenAI

import docx
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter

app = FastAPI(title="NSPU Guide Chatbot Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"]
)

# --- 2. Database & OpenAI Configuration ---
DB_CONFIG = {
    "host": "localhost",
    "user": "root",
    "password": "root",  # <--- Verify your MySQL Password
    "database": "nspu_db",
    "cursorclass": pymysql.cursors.DictCursor
}

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("❌ OPENAI_API_KEY not found! Please check your .env file in the root folder.")

openai_client = OpenAI(api_key=OPENAI_API_KEY)

def get_db_connection():
    return pymysql.connect(**DB_CONFIG)

def init_db_tables():
    """Ensure all required tables exist in MySQL on startup"""
    try:
        connection = get_db_connection()
        with connection.cursor() as cursor:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS comments (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    user_role VARCHAR(50) DEFAULT 'Student',
                    comment_text TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS absa_results (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    comment_id INT NOT NULL,
                    aspect VARCHAR(50) NOT NULL,
                    sentiment VARCHAR(20) NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (comment_id) REFERENCES comments(id) ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS admin_users (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    username VARCHAR(50) UNIQUE,
                    password VARCHAR(255) NOT NULL
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """)
            cursor.execute("INSERT IGNORE INTO admin_users (id, username, password) VALUES (1, 'admin', 'root')")
        connection.commit()
        connection.close()
        print("✅ Database tables verified and ready.")
    except Exception as e:
        print(f"⚠️ MySQL table verification warning: {e}")

init_db_tables()

def verify_admin_password(input_password: str) -> bool:
    if not input_password:
        return False
    # Fallback to standard master passwords
    if input_password in ["root", "admin123"]:
        return True
    try:
        connection = get_db_connection()
        with connection.cursor() as cursor:
            sql = "SELECT id FROM admin_users WHERE password = %s"
            cursor.execute(sql, (input_password,))
            result = cursor.fetchone()
            return result is not None
    except Exception:
        return False
    finally:
        try:
            connection.close()
        except Exception:
            pass

# --- 3. WebSocket Manager ---
class DashboardConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)

    async def broadcast(self, message: dict):
        for connection in list(self.active_connections):
            try:
                await connection.send_json(message)
            except Exception:
                pass

manager = DashboardConnectionManager()

class CommentInput(BaseModel):
    user_role: str = "Student"
    comment_text: str
    manual_aspect: Optional[str] = None
    manual_sentiment: Optional[str] = None

# --- 4. High-Precision OpenAI Embeddings ---
def get_text_embedding(text: str):
    response = openai_client.embeddings.create(
        input=text,
        model="text-embedding-3-small"
    )
    return response.data[0].embedding

def get_text_embeddings_batch(texts: List[str]):
    response = openai_client.embeddings.create(
        input=texts,
        model="text-embedding-3-small"
    )
    return [item.embedding for item in response.data]

# --- 5. Structured Document Extraction ---
def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        extracted = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                extracted.append(txt)
                
        for table in doc.tables:
            table_lines = ["\n[ဇယား အချက်အလက်များ]"]
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                clean_row = []
                for cell_txt in row_cells:
                    if not clean_row or clean_row[-1] != cell_txt:
                        clean_row.append(cell_txt)
                if clean_row:
                    table_lines.append(" - " + " | ".join(clean_row))
            extracted.append("\n".join(table_lines))
                    
        return "\n\n".join(extracted)
    except Exception as e:
        print(f"⚠️ Error reading {file_path}: {e}")
        return ""

def extract_text_from_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
        pdf_text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pdf_text += t + "\n\n"
        return pdf_text
    except Exception as e:
        print(f"⚠️ Error reading {file_path}: {e}")
        return ""

# --- 6. Intelligent Knowledge Base Initialization ---
FULL_DOCUMENT_TEXT = ""
DOCUMENT_CHUNKS = []

def init_vector_db():
    global FULL_DOCUMENT_TEXT, DOCUMENT_CHUNKS
    extracted_text = ""
    processed_files_count = 0
    
    for file in os.listdir("."):
        if file.endswith(".docx") and not file.startswith("~$"):
            print(f"💡 Reading Word Document: {file}...")
            text = extract_text_from_docx(file)
            if text.strip():
                extracted_text += f"\n\n=== SOURCE: {file} ===\n\n" + text
                processed_files_count += 1
                
        elif file.endswith(".pdf"):
            print(f"💡 Reading PDF Document: {file}...")
            text = extract_text_from_pdf(file)
            if text.strip():
                extracted_text += f"\n\n=== SOURCE: {file} ===\n\n" + text
                processed_files_count += 1

    if extracted_text.strip():
        FULL_DOCUMENT_TEXT = extracted_text
        print(f"💡 Indexing {processed_files_count} document(s) with OpenAI Embeddings...")
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=250,
            separators=["\n\n", "\n", "။", "၊", " "]
        )
        DOCUMENT_CHUNKS = text_splitter.split_text(extracted_text)
        
        chroma_client = chromadb.PersistentClient(path="./nspu_vector_db")
        try:
            chroma_client.delete_collection(name="nspu_docs")
        except Exception:
            pass
            
        collection = chroma_client.create_collection(name="nspu_docs")
        embeddings = get_text_embeddings_batch(DOCUMENT_CHUNKS)
        ids = [f"doc_{i}" for i in range(len(DOCUMENT_CHUNKS))]
        
        collection.add(
            documents=DOCUMENT_CHUNKS,
            embeddings=embeddings,
            ids=ids
        )
        print(f"✅ Vector DB Created Successfully! (Total Chunks Indexed: {len(DOCUMENT_CHUNKS)})")
    else:
        print("⚠️ Warning: No valid .docx or .pdf documents found.")

init_vector_db()

# --- 7. API Endpoints ---

@app.post("/api/comments")
@app.post("/api/feedback")
async def create_comment(data: CommentInput):
    connection = get_db_connection()
    try:
        if data.manual_aspect and data.manual_sentiment:
            aspect = data.manual_aspect
            sentiment = data.manual_sentiment
        else:
            prompt = f"""
            You are analyzing student feedback for a university in Myanmar.
            Categorize the comment into EXACTLY ONE aspect:
            - "Environment" (ကျောင်းပတ်ဝန်းကျင်၊ သန့်ရှင်းရေး၊ အမှိုက်ပုံး၊ သစ်ပင်၊ လေကောင်းလေသန့်၊ သာယာလှပမှု)
            - "Facilities" (စာသင်ခန်း၊ အဲကွန်း/ပန်ကာ၊ Wi-Fi၊ Lab ခန်း၊ စာကြည့်တိုက်၊ အားကစားရုံ၊ မီး/ရေ)
            - "Teaching" (ဆရာ/ဆရာမ သင်ကြားပြသမှု၊ သင်ရိုးညွှန်းတမ်း၊ စာမေးပွဲ၊ လက်တွေ့သင်ကြားမှု)
            - "Administrative" (ကျောင်းအပ်နှံခြင်း၊ ရုံးလုပ်ငန်း၊ ဝန်ထမ်းများဆက်ဆံရေး၊ စည်းကမ်းထိန်းသိမ်းရေး)
            - "Canteen" (ကျောင်းမုန့်ဈေးတန်း၊ အစားအသောက် အရည်အသွေး၊ ဈေးနှုန်း၊ ကန်တင်းသန့်ရှင်းရေး)

            Determine sentiment: "Positive", "Negative", or "Neutral".

            Comment: "{data.comment_text}"
            Return strictly JSON: {{"aspect": "...", "sentiment": "..."}}
            """
            response = openai_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="gpt-4o-mini",
                response_format={"type": "json_object"}
            )
            analysis = json.loads(response.choices[0].message.content)
            aspect = analysis.get("aspect", "Administrative")
            sentiment = analysis.get("sentiment", "Neutral")

        # Normalize Aspect & Sentiment
        if aspect in ["Facility", "Facilities"]:
            aspect = "Facilities"

        if sentiment in ["Good", "Positive"]:
            sentiment = "Positive"
        elif sentiment in ["Bad", "Negative"]:
            sentiment = "Negative"
        else:
            sentiment = "Neutral"

        with connection.cursor() as cursor:
            sql_comment = "INSERT INTO comments (user_role, comment_text) VALUES (%s, %s)"
            cursor.execute(sql_comment, (data.user_role, data.comment_text))
            comment_id = cursor.lastrowid
            
            sql_absa = "INSERT INTO absa_results (comment_id, aspect, sentiment) VALUES (%s, %s, %s)"
            cursor.execute(sql_absa, (comment_id, aspect, sentiment))
            
        connection.commit()
        
        feedback_data = {
            "id": comment_id,
            "user_role": data.user_role,
            "role": data.user_role,
            "comment_text": data.comment_text,
            "aspect": aspect,
            "sentiment": sentiment
        }
        
        await manager.broadcast({
            "event": "new_feedback_received",
            "data": feedback_data
        })
        
        return {"status": "success", "result": feedback_data}
    except Exception as e:
        connection.rollback()
        print(f"Error saving feedback: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        connection.close()


@app.get("/api/ask")
async def ask_chatbot(question: str = Query(..., description="User Question")):
    clean_q = question.strip()
    if not clean_q:
        return {"answer": "မင်္ဂလာပါခင်ဗျာ၊ သိရှိလိုသည်များကို မေးမြန်းနိုင်ပါသည်။"}
    
    if not DOCUMENT_CHUNKS or not FULL_DOCUMENT_TEXT:
        raise HTTPException(status_code=400, detail="Knowledge base not ready.")

    try:
        # Step 1: AI Query Expansion (Generates semantic variations & key concepts)
        expand_prompt = f"""
        User Query: "{clean_q}"
        Task: Identify the core topic and generate 3 to 5 related Myanmar synonyms, English terms, and topic keywords.
        Example: "ကျင်းပသောပွဲအခမ်းအနားများ" -> ["ပွဲ", "အခမ်းအနား", "Fresher", "Festival", "Dinner", "လှုပ်ရှားမှု"]
        Return strictly a JSON list of strings: {{"keywords": ["term1", "term2", ...]}}
        """
        try:
            exp_resp = openai_client.chat.completions.create(
                messages=[{"role": "user", "content": expand_prompt}],
                model="gpt-4o-mini",
                response_format={"type": "json_object"}
            )
            parsed = json.loads(exp_resp.choices[0].message.content)
            search_terms = parsed.get("keywords", [clean_q])
        except Exception:
            search_terms = [clean_q]

        # Ensure original query is always included
        search_terms.append(clean_q)

        # Step 2: Semantic Vector Search
        chroma_client = chromadb.PersistentClient(path="./nspu_vector_db")
        collection = chroma_client.get_collection(name="nspu_docs")
        q_embedding = get_text_embedding(clean_q)
        
        vec_results = collection.query(
            query_embeddings=[q_embedding],
            n_results=6
        )
        matched_chunks = vec_results["documents"][0] if vec_results.get("documents") else []

        # Step 3: Fuzzy Substring Search across all chunks using expanded terms
        keyword_chunks = []
        for chunk in DOCUMENT_CHUNKS:
            for term in search_terms:
                if isinstance(term, str) and len(term.strip()) >= 2 and term.strip().lower() in chunk.lower():
                    keyword_chunks.append(chunk)
                    break

        # Step 4: Combine & Deduplicate Contexts
        all_matched = []
        for c in (keyword_chunks + matched_chunks):
            if c not in all_matched:
                all_matched.append(c)

        # Fallback: If matches are extremely low, provide the most relevant top sections
        if len(all_matched) < 2:
            all_matched = DOCUMENT_CHUNKS[:4]

        context = "\n\n---\n\n".join(all_matched[:8])

        # Step 5: Flexible System Prompt for Comprehensive Answering
        system_prompt = """You are the knowledgeable and polite student guide assistant for Naypyitaw State Polytechnic University (NSPU).
Use the provided Context to answer the user's question completely in Myanmar Unicode (Padauk style).

Guidelines:
1. If the user asks about events, rules, faculties, admission criteria, canteen, or facilities, extract and summarize ALL relevant points found in the Context. Use bullet points for readability.
2. Synthesize context intelligently. For example, if they ask for "ပွဲအခမ်းအနားများ", list "Fresher Welcome", "Science Festival", etc. even if the exact phrase "ပွဲအခမ်းအနားများ" isn't explicitly written next to those events.
3. Be helpful and clear. 
4. Only if the Context completely lacks any related information on the topic, politely inform the user that the guide does not contain that specific detail."""

        user_content = f"Context Information from University Documents:\n{context}\n\nUser Question: {clean_q}"

        response = openai_client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            model="gpt-4o-mini",
            temperature=0.1
        )

        return {"answer": response.choices[0].message.content}

    except Exception as e:
        print(f"Error answering question: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/admin/comments")
async def get_all_comments(x_admin_secret: str = Header(None)):
    if not verify_admin_password(x_admin_secret):
        raise HTTPException(status_code=401, detail="Unauthorized Access")
        
    connection = get_db_connection()
    try:
        with connection.cursor() as cursor:
            sql = """
                SELECT c.id, c.user_role, c.comment_text, 
                       DATE_FORMAT(c.created_at, '%Y-%m-%d %H:%i:%s') AS created_at, 
                       a.aspect, a.sentiment 
                FROM comments c
                LEFT JOIN absa_results a ON c.id = a.comment_id
                ORDER BY c.id DESC
            """
            cursor.execute(sql)
            results = cursor.fetchall()
            return {"status": "success", "data": results}
    except Exception as e:
        print(f"Error fetching admin comments: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        connection.close()

@app.delete("/api/admin/comments/{comment_id}")
async def delete_comment(comment_id: int, x_admin_secret: str = Header(None)):
    if not verify_admin_password(x_admin_secret):
        raise HTTPException(status_code=401, detail="Unauthorized Access")
        
    connection = get_db_connection()
    try:
        with connection.cursor() as cursor:
            cursor.execute("DELETE FROM absa_results WHERE comment_id = %s", (comment_id,))
            cursor.execute("DELETE FROM comments WHERE id = %s", (comment_id,))
        connection.commit()
        return {"status": "success", "message": f"Comment {comment_id} deleted."}
    except Exception as e:
        connection.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        connection.close()

@app.websocket("/ws/dashboard")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        manager.disconnect(websocket)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)